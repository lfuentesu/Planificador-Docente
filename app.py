import io
import os
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from dotenv import load_dotenv
from groq import Groq

# Cargar variables de entorno si existen localmente
load_dotenv()

# Intentar importar Supabase para la base de datos comunitaria
try:
    from supabase import create_client, Client
    SUPABASE_DISPONIBLE = True
except ImportError:
    SUPABASE_DISPONIBLE = False

# Importar el módulo de legislación laboral desde la carpeta
try:
    from legislacion.pagina_legislacion import mostrar_seccion_legislacion
    LEGISLACION_DISPONIBLE = True
except ImportError:
    LEGISLACION_DISPONIBLE = False

# Configuración inicial de la página
st.set_page_config(
    page_title="Planificador y Asistente Docente", page_icon="📚", layout="wide"
)

# --- INYECCIÓN DE ESTILOS CSS PERSONALIZADOS ---
st.markdown(
    """
    <style>
    /* Estilo general del fondo */
    .stApp {
        background-color: #F8FAFC;
    }

    /* Encabezados principales */
    h1 {
        color: #1E3A8A !important;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        font-weight: 700 !important;
    }
    
    h2, h3 {
        color: #1E40AF !important;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }

    /* Tarjetas estilizadas para secciones */
    div.stExpander, div[data-testid="stForm"] {
        background-color: #FFFFFF;
        border-radius: 12px;
        border: 1px solid #E2E8F0;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05);
        padding: 10px;
    }

    /* Estilo para la barra lateral */
    section[data-testid="stSidebar"] {
        background-color: #0F172A;
    }

    section[data-testid="stSidebar"] * {
        color: #F8FAFC !important;
    }

    section[data-testid="stSidebar"] .stRadio label {
        color: #E2E8F0 !important;
        font-weight: 500;
    }

    /* Botones primarios (Generar / Descargar / Publicar) */
    div.stButton > button[kind="primary"] {
        background-color: #2563EB !important;
        color: white !important;
        border-radius: 8px !important;
        border: none !important;
        font-weight: 600 !important;
        padding: 0.6rem 1.2rem !important;
        transition: all 0.3s ease;
        box-shadow: 0 2px 4px rgba(37, 99, 235, 0.2);
    }

    div.stButton > button[kind="primary"]:hover {
        background-color: #1D4ED8 !important;
        box-shadow: 0 4px 8px rgba(29, 78, 216, 0.3);
        transform: translateY(-1px);
    }

    /* Botón de descarga en Word */
    div.stDownloadButton > button {
        background-color: #059669 !important;
        color: white !important;
        border-radius: 8px !important;
        border: none !important;
        font-weight: 600 !important;
        padding: 0.6rem 1.2rem !important;
        transition: all 0.3s ease;
    }

    div.stDownloadButton > button:hover {
        background-color: #047857 !important;
        transform: translateY(-1px);
    }

    /* Tarjeta de solidaridad docente */
    .solidaridad-card {
        background-color: #EFF6FF;
        border-left: 5px solid #2563EB;
        padding: 15px 20px;
        border-radius: 8px;
        margin-top: 15px;
        margin-bottom: 15px;
    }

    /* Cajas de texto y campos de selección */
    .stTextInput>div>div>input, .stSelectbox>div>div>div, .stTextArea>div>div>textarea {
        border-radius: 8px !important;
        border: 1px solid #CBD5E1 !important;
    }

    .stAlert {
        border-radius: 10px !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# Diccionario de ejemplos automáticos según la asignatura seleccionada
EJEMPLOS_OBJETIVOS = {
    "Matemáticas": "Suma y resta de fracciones con distinto denominador",
    "Lengua y Literatura / Lenguaje": "Comprensión lectora y análisis de personajes en mitos grecolatinos",
    "Historia, Geografía y Ciencias Sociales": "Causas y consecuencias del proceso de Independencia de Chile",
    "Ciencias Naturales (Biología, Física, Química)": "Estructura de la célula eucariota y función de sus organelos",
    "Inglés": "Uso del presente simple y continuo para describir rutinas diarias",
    "Artes Visuales / Música": "Elementos del lenguaje visual: color, línea y textura en obras impresionistas",
    "Tecnología": "Diseño de un prototipo de solución a un problema medioambiental local",
    "Educación Física y Salud": "Desarrollo de cualidades físicas básicas mediante juegos colectivos",
    "Orientación": "Promoción del respeto y la resolución pacífica de conflictos en la comunidad escolar",
}

def inicializar_supabase():
    """Inicializa la conexión con Supabase si las credenciales existen en Secrets."""
    url = st.secrets.get("SUPABASE_URL", os.getenv("SUPABASE_URL", ""))
    key = st.secrets.get("SUPABASE_KEY", os.getenv("SUPABASE_KEY", ""))
    if url and key and SUPABASE_DISPONIBLE:
        try:
            return create_client(url, key)
        except Exception:
            return None
    return None

def generar_planificacion_ia(groq_api_key, nivel, asignatura, tipo_planificacion, enfoque, observaciones):
    """Llama a la API de Groq para generar una planificación pedagógica detallada."""
    client = Groq(api_key=groq_api_key)

    prompt_sistema = """
    Eres un experto diseñador curricular y asesor pedagógico del sistema educativo chileno, con amplio conocimiento del Marco para la Buena Enseñanza, las Bases Curriculares del MINEDUC y los Decretos de Evaluación (ej. Decreto 67).

    Tu objetivo es redactar planificaciones pedagógicas altamente estructuradas, rigurosas, claras y listas para ser implementadas por los docentes en el aula.

    Asegúrate de incluir siempre las siguientes secciones bien detalladas:

    1. **Objetivo de Aprendizaje (OA)** y/o Habilidad principal adaptado al curso y asignatura específicos.
    2. **Indicadores de Evaluación** específicos y medibles.
    3. **Estructura de la Clase:**
       - **Inicio (15 min):** Activación de conocimientos previos, conflicto cognitivo y declaración del objetivo.
       - **Desarrollo (60 min):** Modelaje del docente, práctica guiada y práctica independiente con actividades concretas.
       - **Cierre (15 min):** Síntesis de la clase.
    4. **🎫 Ticket de Salida:** Una pregunta, ejercicio práctico o reflexión breve y concreta para aplicar en los últimos 5 minutos y verificar el aprendizaje formativo.
    5. **📊 Instrumento de Evaluación Sugerido:** Incluye una propuesta de Pauta de Cotejo, Rúbrica Breve o Escala de Apreciación con 3 a 4 criterios para evaluar la clase/unidad.
    6. **Sugerencia de Diversificación (DUA):** Adecuaciones concretas para atender a la diversidad en el aula.
    7. **Recursos Pedagógicos Sugeridos.**

    Usa un formato Markdown limpio, profesional y bien organizado con encabezados y viñetas.
    """

    prompt_usuario = f"""
    Genera una propuesta de planificación con los siguientes parámetros:
    - **Nivel / Curso Específico:** {nivel}
    - **Asignatura:** {asignatura}
    - **Tipo de Planificación:** {tipo_planificacion}
    - **Tema / Objetivo Específico:** {enfoque}
    - **Consideraciones adicionales:** {observaciones if observaciones else "Ninguna especificada"}
    """

    respuesta = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": prompt_sistema},
            {"role": "user", "content": prompt_usuario},
        ],
        temperature=0.4,
        max_tokens=2500,
    )

    return respuesta.choices[0].message.content

def crear_documento_word(nivel, asignatura, tipo_planificacion, enfoque, observaciones, contenido):
    """Crea un archivo Word (.docx) formateado con la planificación."""
    doc = Document()

    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run("PROPUESTA DE PLANIFICACIÓN PEDAGÓGICA")
    run_titulo.bold = True
    run_titulo.font.size = Pt(16)
    run_titulo.font.color.rgb = RGBColor(0, 51, 102)
    p_titulo.paragraph_format.space_after = Pt(12)

    tabla = doc.add_table(rows=4, cols=2)
    tabla.style = "Table Grid"

    datos = [
        ("Nivel / Curso:", nivel),
        ("Asignatura:", asignatura),
        ("Tipo de Planificación:", tipo_planificacion),
        ("Tema / Enfoque:", enfoque),
    ]

    for i, (campo, valor) in enumerate(datos):
        cell_label = tabla.cell(i, 0)
        cell_value = tabla.cell(i, 1)

        p_label = cell_label.paragraphs[0]
        run_l = p_label.add_run(campo)
        run_l.bold = True
        run_l.font.size = Pt(10)

        p_val = cell_value.paragraphs[0]
        run_v = p_val.add_run(valor)
        run_v.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    if observaciones.strip():
        p_obs = doc.add_paragraph()
        run_o_title = p_obs.add_run("Consideraciones especiales: ")
        run_o_title.bold = True
        p_obs.add_run(observaciones)
        p_obs.paragraph_format.space_after = Pt(12)

    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run("Desarrollo de la Planificación")
    run_sub.bold = True
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(0, 51, 102)
    p_sub.paragraph_format.space_after = Pt(8)

    for linea in contenido.split("\n"):
        linea_str = linea.strip()
        if not linea_str:
            continue

        p = doc.add_paragraph()
        if linea_str.startswith("# "):
            run = p.add_run(linea_str.replace("# ", ""))
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 51, 102)
        elif linea_str.startswith("## "):
            run = p.add_run(linea_str.replace("## ", ""))
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 51, 102)
        elif linea_str.startswith("### "):
            run = p.add_run(linea_str.replace("### ", ""))
            run.bold = True
            run.font.size = Pt(11)
        elif linea_str.startswith("- ") or linea_str.startswith("* "):
            p.style = "List Bullet"
            texto_limpio = linea_str[2:].replace("**", "")
            p.add_run(texto_limpio)
        else:
            texto_limpio = linea_str.replace("**", "")
            p.add_run(texto_limpio)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def mostrar_seccion_planificaciones(groq_api_key, supabase_client):
    """Módulo del Generador de Planificaciones de Clase impulsado por IA."""
    st.title("📝 Generador de Planificaciones de Clase (IA)")
    st.write("Diseña planificaciones pedagógicas completas alineadas con los estándares del MINEDUC usando Inteligencia Artificial.")

    with st.expander("💡 **¿Cómo utilizar este generador? (Instructivo paso a paso)**", expanded=False):
        st.markdown(
            """
            1. **Selecciona el Nivel y la Asignatura:** Elige el curso exacto (desde Prekínder a 4° Medio/EPJA) y el área de aprendizaje.
            2. **Ingresa el Tema u Objetivo:** Escribe el contenido específico que deseas trabajar en la clase.
            3. **Agrega Consideraciones (Opcional):** Anota detalles DUA, estudiantes en PIE o personalizaciones.
            4. **Genera, Descarga y Comparte:** Al crear la clase podrás descargarla en Word (.docx) y compartirla en la **Biblioteca Comunitaria**.
            """
        )

    st.markdown("---")

    col1, col2 = st.columns(2)
    with col1:
        nivel = st.selectbox(
            "Selecciona el Nivel / Curso:",
            [
                "Prekínder (NT1)", "Kínder (NT2)",
                "1° Básico", "2° Básico", "3° Básico", "4° Básico", "5° Básico", "6° Básico", "7° Básico", "8° Básico",
                "1° Medio", "2° Medio", "3° Medio (Formación General)", "3° Medio (TTP / Técnico Profesional)", "4° Medio (Formación General)", "4° Medio (TTP / Técnico Profesional)",
                "EPJA / Educación de Adultos",
            ],
        )
        asignatura = st.selectbox(
            "Selecciona la Asignatura:",
            [
                "Matemáticas", "Lengua y Literatura / Lenguaje", "Historia, Geografía y Ciencias Sociales",
                "Ciencias Naturales (Biología, Física, Química)", "Inglés", "Artes Visuales / Música",
                "Tecnología", "Educación Física y Salud", "Orientación",
            ],
        )

    ejemplo_sugerido = EJEMPLOS_OBJETIVOS.get(asignatura, "Escribe aquí el objetivo...")

    with col2:
        tipo_planificacion = st.selectbox(
            "Tipo de Planificación:",
            [
                "Clase a Clase (90 minutos)",
                "Unidad Didáctica (2 a 4 semanas)",
                "Planificación Anual (Visión General)",
            ],
        )
        enfoque = st.text_input(
            "Objetivo o Contenido Específico:",
            placeholder=f"Ej: {ejemplo_sugerido}",
            help=f"Sugerencia para {asignatura}: {ejemplo_sugerido}",
        )

    observaciones = st.text_area(
        "Contexto o necesidades específicas (Opcional):",
        placeholder="Ej: Incluir trabajo colaborativo en parejas, priorizar evaluación formativa, considerar estudiantes con DUA...",
        height=80,
    )

    st.markdown("---")

    if st.button("🚀 Generar Planificación con IA", type="primary"):
        if not groq_api_key:
            st.error("⚠️ No se detectó la clave de Groq API. Configura `GROQ_API_KEY` en los Secrets de Streamlit.")
        elif not enfoque.strip():
            st.warning(f"Por favor ingresa un objetivo o contenido para la asignatura de **{asignatura}**.")
        else:
            with st.spinner(f"Generando propuesta pedagógica de {asignatura} para {nivel}... Esto tomará unos segundos."):
                try:
                    resultado_planificacion = generar_planificacion_ia(
                        groq_api_key, nivel, asignatura, tipo_planificacion, enfoque, observaciones
                    )

                    st.session_state["resultado_planificacion"] = resultado_planificacion
                    st.session_state["datos_planificacion"] = {
                        "nivel": nivel,
                        "asignatura": asignatura,
                        "tipo": tipo_planificacion,
                        "enfoque": enfoque,
                        "observaciones": observaciones,
                    }
                    st.success("¡Planificación generada con éxito!")

                except Exception as e:
                    st.error(f"Ocurrió un error al generar la planificación mediante IA: {e}")

    if "resultado_planificacion" in st.session_state:
        st.markdown("### 📄 Propuesta Pedagógica")
        
        st.markdown(
            f'<div style="background-color: #FFFFFF; padding: 25px; border-radius: 12px; border: 1px solid #E2E8F0; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05); margin-bottom: 20px;">{st.session_state["resultado_planificacion"]}</div>',
            unsafe_allow_html=True,
        )

        datos = st.session_state["datos_planificacion"]

        col_down, col_pub = st.columns([1, 1])

        with col_down:
            archivo_word = crear_documento_word(
                datos["nivel"], datos["asignatura"], datos["tipo"], datos["enfoque"], datos["observaciones"], st.session_state["resultado_planificacion"]
            )
            nombre_archivo = f"Planificacion_{datos['asignatura']}_{datos['nivel']}.docx".replace(" ", "_")

            st.download_button(
                label="📥 Descargar Planificación en Word (.docx)",
                data=archivo_word,
                file_name=nombre_archivo,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        with col_pub:
            if st.button("🤝 Publicar en la Biblioteca Comunitaria", type="primary"):
                if supabase_client:
                    try:
                        supabase_client.table("planificaciones").insert({
                            "nivel": datos["nivel"],
                            "asignatura": datos["asignatura"],
                            "tipo": datos["tipo"],
                            "enfoque": datos["enfoque"],
                            "contenido": st.session_state["resultado_planificacion"]
                        }).execute()
                        st.balloons()
                        st.success("¡Muchas gracias! Tu planificación ha sido compartida con la comunidad docente.")
                    except Exception as e:
                        st.error(f"Error al guardar en la biblioteca: {e}")
                else:
                    if "biblioteca_local" not in st.session_state:
                        st.session_state["biblioteca_local"] = []
                    st.session_state["biblioteca_local"].append({
                        "nivel": datos["nivel"],
                        "asignatura": datos["asignatura"],
                        "tipo": datos["tipo"],
                        "enfoque": datos["enfoque"],
                        "contenido": st.session_state["resultado_planificacion"]
                    })
                    st.balloons()
                    st.success("¡Gracias por colaborar! La planificación quedó registrada para la comunidad.")

        st.markdown(
            """
            <div class="solidaridad-card">
                <h4 style="margin: 0; color: #1E40AF; font-size: 1.1rem;">🤝 <b>Espacio de Solidaridad y Colaboración Mutua</b></h4>
                <p style="margin: 5px 0 0 0; color: #334155; font-size: 0.95rem;">
                    Al compartir tus propuestas en la <b>Biblioteca Comunitaria</b>, estás solidarizando activamente con colegas de todo el país. Juntos hacemos de esta herramienta un espacio colectivo para aliviar la carga administrativa y enriquecer la práctica pedagógica en Chile.
                </p>
            </div>
            """,
            unsafe_allow_html=True,
        )

def mostrar_seccion_biblioteca(supabase_client):
    """Módulo de consulta de la Biblioteca Comunitaria de Planificaciones."""
    st.title("📚 Biblioteca Comunitaria de Planificaciones")
    st.write("Explora y reutiliza planificaciones pedagógicas compartidas de forma solidaria por otros docentes.")
    st.markdown("---")

    col_f1, col_f2 = st.columns(2)
    with col_f1:
        filtro_asig = st.selectbox("Filtrar por Asignatura:", ["Todas"] + list(EJEMPLOS_OBJETIVOS.keys()))
    with col_f2:
        filtro_busqueda = st.text_input("Buscar por palabra clave:", placeholder="Ej: fracciones, independencia, celular...")

    lista_planificaciones = []
    if supabase_client:
        try:
            res = supabase_client.table("planificaciones").select("*").order("created_at", desc=True).execute()
            lista_planificaciones = res.data
        except Exception:
            lista_planificaciones = st.session_state.get("biblioteca_local", [])
    else:
        lista_planificaciones = st.session_state.get("biblioteca_local", [])

    resultados = []
    for item in lista_planificaciones:
        cumple_asig = (filtro_asig == "Todas") or (item.get("asignatura") == filtro_asig)
        cumple_txt = True
        if filtro_busqueda.strip():
            kw = filtro_busqueda.lower()
            cumple_txt = (kw in item.get("enfoque", "").lower()) or (kw in item.get("contenido", "").lower())
        
        if cumple_asig and cumple_txt:
            resultados.append(item)

    if not resultados:
        st.info("ℹ️ Aún no hay planificaciones registradas en este filtro. ¡Sé el primero en aportar a la comunidad!")
    else:
        st.success(f"Se encontraron {len(resultados)} planificaciones comunitarias:")
        for idx, item in enumerate(resultados):
            with st.expander(f"📌 **{item.get('asignatura')}** — {item.get('nivel')} | *{item.get('enfoque')}*"):
                st.markdown(f"**Tipo:** {item.get('tipo')}")
                st.markdown("---")
                st.markdown(item.get("contenido"))
                
                buf = crear_documento_word(
                    item.get('nivel', 'N/A'), item.get('asignatura', 'N/A'), item.get('tipo', 'N/A'), item.get('enfoque', 'N/A'), '', item.get('contenido', '')
                )
                st.download_button(
                    label="📥 Descargar esta clase compartida en Word (.docx)",
                    data=buf,
                    file_name=f"Comunitaria_{item.get('asignatura')}_{item.get('nivel')}.docx".replace(" ", "_"),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_{idx}"
                )

def main():
    groq_api_key = st.secrets.get("GROQ_API_KEY", os.getenv("GROQ_API_KEY", ""))
    supabase_client = inicializar_supabase()

    # --- BÚSQUEDA DEL LOGO Y CONTROL DE TAMAÑO EN LA BARRA LATERAL ---
    posibles_rutas = [
        "logotiposute.jpg", "logotiposute.JPG", "logotiposute.jpeg", "logotiposute.png",
        "logotposute.jpg", "legislacion/logotiposute.jpg", "assets/logotiposute.jpg",
    ]

    logo_encontrado = False
    for ruta in posibles_rutas:
        if os.path.exists(ruta):
            # Fijamos width=180 para evitar que la imagen ocupe todo el espacio vertical
            st.sidebar.image(ruta, width=180)
            logo_encontrado = True
            break

    if not logo_encontrado:
        for raiz, carpetas, archivos in os.walk("."):
            for archivo in archivos:
                if "sute" in archivo.lower() and archivo.lower().endswith((".jpg", ".jpeg", ".png")):
                    st.sidebar.image(os.path.join(raiz, archivo), width=180)
                    logo_encontrado = True
                    break
            if logo_encontrado:
                break

    st.sidebar.title("📌 Navegación")
    st.sidebar.markdown("---")

    opcion = st.sidebar.radio(
        "Ir a:",
        [
            "📝 Generador de Planificaciones",
            "📚 Biblioteca Comunitaria",
            "⚖️ Asistente de Legislación Laboral",
        ],
    )

    st.sidebar.markdown("---")
    st.sidebar.caption("🤖 **Planificador Docente v2.2**")
    st.sidebar.caption("Chile — Red Solidaria & Normativa Laboral")

    if opcion == "📝 Generador de Planificaciones":
        mostrar_seccion_planificaciones(groq_api_key, supabase_client)

    elif opcion == "📚 Biblioteca Comunitaria":
        mostrar_seccion_biblioteca(supabase_client)

    elif opcion == "⚖️ Asistente de Legislación Laboral":
        if LEGISLACION_DISPONIBLE:
            mostrar_seccion_legislacion()
        else:
            st.error("⚠️ No se encontró el módulo `legislacion/pagina_legislacion.py`.")

if __name__ == "__main__":
    main()
